"""
docx_builder.py — Python DOCX report generator (replaces generateDocx() in JS).

Translates the frontend generateDocx() + buildHeaderTable() functions to python-docx.
Called by server.py /generate endpoint.
"""
import base64
import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

from PIL import Image


# ─── CONSTANTS ────────────────────────────────────────────────────────────────

BLUE = RGBColor(0x0F, 0x4C, 0x75)
TEAL = RGBColor(0x1B, 0x9A, 0xAA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x99, 0x99, 0x99)

DB_COLORS = [
    (0x0F, 0x4C, 0x75),
    (0x06, 0x7E, 0x7B),
    (0x06, 0x7E, 0x3A),
    (0x6A, 0x1B, 0x9A),
    (0xFF, 0x94, 0x65),
    (0xE6, 0x51, 0x00),
    (0x00, 0x69, 0x5C),
    (0x37, 0x47, 0x4F),
]

PAGE_W_CM = 21.0   # A4 width
MARGIN_CM = 2.54   # 1 inch margins


# ─── HELPERS ──────────────────────────────────────────────────────────────────

def b64_to_bytes(data_url: str) -> Optional[bytes]:
    """Strip data URL prefix and decode base64 to bytes."""
    if not data_url:
        return None
    try:
        if "," in data_url:
            data_url = data_url.split(",", 1)[1]
        return base64.b64decode(data_url)
    except Exception:
        return None


def hex_to_rgb(hex_str: str) -> tuple:
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, rgb: tuple):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb))
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "FFFFFF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_row_height(row, twips: int):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def set_col_span(cell, span: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    gridSpan = OxmlElement("w:gridSpan")
    gridSpan.set(qn("w:val"), str(span))
    tcPr.append(gridSpan)


def set_row_span_start(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement("w:vMerge")
    vMerge.set(qn("w:val"), "restart")
    tcPr.append(vMerge)


def set_row_span_continue(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement("w:vMerge")
    tcPr.append(vMerge)


def set_cell_valign(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def add_run(para, text: str, bold=False, color=None, size_pt=10, font="Cambria"):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def hdr_cell(table_cell, text: str, rgb: tuple = None, size_pt=10, bold=True, font="Cambria"):
    """Style a cell as a header (blue bg, white text)."""
    if rgb is None:
        rgb = (0x0F, 0x4C, 0x75)
    set_cell_bg(table_cell, rgb)
    set_cell_border(table_cell)
    set_cell_margins(table_cell)
    para = table_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = WHITE
    set_cell_valign(table_cell)
    return table_cell


def data_cell(table_cell, text: str, size_pt=9, bold=False, font="Cambria", bg_rgb=None):
    """Style a cell as a data cell."""
    if bg_rgb:
        set_cell_bg(table_cell, bg_rgb)
    set_cell_border(table_cell)
    set_cell_margins(table_cell, top=40, bottom=40, left=60, right=60)
    para = table_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    safe_text = wrap_long(str(text or ""))
    run = para.add_run(safe_text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    set_cell_valign(table_cell)
    return table_cell


def wrap_long(text: str, max_len=17) -> str:
    """Insert zero-width spaces in words longer than max_len."""
    parts = []
    for word in text.split(" "):
        if len(word) > max_len:
            chunks = re.findall(f".{{1,{max_len}}}", word)
            parts.append("​".join(chunks))
        else:
            parts.append(word)
    return " ".join(parts)


def set_table_width(table, width_cm: float):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))  # 1 cm ≈ 567 twips
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def docx_page_break():
    """Return an XML page break element to append to a run's _r."""
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def add_page_break_para(doc):
    """Add a paragraph containing a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run._r.append(docx_page_break())
    return para


def insert_image(doc, img_bytes: bytes, width_inches=5.0) -> bool:
    """Add an image paragraph to the document. Returns True on success."""
    try:
        stream = io.BytesIO(img_bytes)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(stream, width=Inches(width_inches))
        return True
    except Exception:
        return False


def fit_image(img_bytes: bytes, max_w_inches: float, max_h_inches: float) -> tuple:
    """Return (width, height) in Inches that fit within max bounds, preserving aspect ratio.
    Only scales DOWN if too large — never upscales a small image.
    Falls back to (max_w_inches, None) if Pillow can't read the image."""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        orig_w_px, orig_h_px = img.size
        if orig_w_px == 0 or orig_h_px == 0:
            return Inches(max_w_inches), None
        # Convert px to inches at 96 dpi (screen resolution default)
        dpi = img.info.get("dpi", (96, 96))
        dpi_x = dpi[0] if dpi[0] > 0 else 96
        dpi_y = dpi[1] if dpi[1] > 0 else 96
        natural_w = orig_w_px / dpi_x
        natural_h = orig_h_px / dpi_y
        ratio = natural_w / natural_h
        # Start with natural size, then clamp to max box
        w = min(natural_w, max_w_inches)
        h = w / ratio
        if h > max_h_inches:
            h = max_h_inches
            w = h * ratio
        return Inches(w), Inches(h)
    except Exception:
        return Inches(max_w_inches), None


def add_logo_to_run(run, img_bytes: bytes, max_w_inches: float, max_h_inches: float):
    """Add a proportionally scaled logo picture to an existing run."""
    w, h = fit_image(img_bytes, max_w_inches, max_h_inches)
    kwargs = {"width": w}
    if h is not None:
        kwargs["height"] = h
    run.add_picture(io.BytesIO(img_bytes), **kwargs)


# ─── CPU CHART ────────────────────────────────────────────────────────────────

def render_cpu_chart(cpu_data: dict) -> Optional[bytes]:
    """Render CPU utilization chart to PNG bytes using matplotlib."""
    rows = cpu_data.get("rows", [])
    cols = cpu_data.get("columns", [])
    if not rows or not cols:
        return None

    def col_idx(pattern):
        p = re.compile(pattern, re.IGNORECASE)
        for i, c in enumerate(cols):
            if p.search(str(c)):
                return i
        return -1

    l_idx = col_idx(r"begin|snap")
    if l_idx < 0:
        l_idx = 0
    m_idx = col_idx(r"max")
    a_idx = col_idx(r"avg")
    inst_idx = col_idx(r"instance")

    def to_n(v):
        try:
            return float(v)
        except (ValueError, TypeError):
            return None

    instances = []
    if inst_idx >= 0:
        seen = set()
        for r in rows:
            if inst_idx < len(r) and r[inst_idx] and r[inst_idx] not in seen:
                seen.add(r[inst_idx])
                instances.append(r[inst_idx])

    COLORS = ["#d44444", "#1b9aaa", "#f59e0b", "#8b5cf6", "#22c55e", "#f97316"]

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.set_ylim(0, 100)
    ax.set_ylabel("CPU %", fontsize=9)
    ax.grid(color="#e8f1f5", linestyle="-", linewidth=0.5)

    if len(instances) > 1:
        for i, inst in enumerate(instances):
            inst_rows = [r for r in rows if inst_idx < len(r) and r[inst_idx] == inst]
            labels = [r[l_idx][:16] if l_idx < len(r) else "" for r in inst_rows]
            maxv = [to_n(r[m_idx]) if m_idx >= 0 and m_idx < len(r) else None for r in inst_rows]
            ax.plot(labels, maxv, label=f"Inst {inst} Max%",
                    color=COLORS[i % len(COLORS)], linewidth=1.5, marker="o", markersize=2)
    else:
        labels = [r[l_idx][:16] if l_idx < len(r) else "" for r in rows]
        maxv = [to_n(r[m_idx]) if m_idx >= 0 and m_idx < len(r) else None for r in rows]
        avgv = [to_n(r[a_idx]) if a_idx >= 0 and a_idx < len(r) else None for r in rows]
        ax.plot(labels, maxv, label="Max CPU%", color="#d44444", linewidth=1.5, marker="o", markersize=2)
        if a_idx >= 0:
            ax.plot(labels, avgv, label="Avg CPU%", color="#1b9aaa", linewidth=1.5, marker="o", markersize=2)

    tick_step = max(1, len(rows) // 12)
    ax.set_xticks(range(0, len(rows), tick_step))
    xlabels = [rows[i][l_idx][:16] if l_idx < len(rows[i]) else "" for i in range(0, len(rows), tick_step)]
    ax.set_xticklabels(xlabels, rotation=45, fontsize=7, ha="right")
    ax.legend(fontsize=8, loc="upper right")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─── DOCUMENT STYLES ──────────────────────────────────────────────────────────

def apply_doc_styles(doc: Document):
    """Set default font and heading styles, creating them if missing."""
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(10)

    for h_name, h_id, size in [
        ("Heading 1", "Heading1", 14),
        ("Heading 2", "Heading2", 12),
        ("Heading 3", "Heading3", 11),
    ]:
        try:
            h = doc.styles[h_name]
        except KeyError:
            h = doc.styles.add_style(h_name, WD_STYLE_TYPE.PARAGRAPH)
            h.base_style = doc.styles["Normal"]
            pf = h.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            # Set outline level so TOC picks it up
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pPr = h.element.get_or_add_pPr()
            outlineLvl = OxmlElement("w:outlineLvl")
            level = {"Heading 1": "0", "Heading 2": "1", "Heading 3": "2"}[h_name]
            outlineLvl.set(qn("w:val"), level)
            pPr.append(outlineLvl)
        h.font.name = "Cambria"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = BLACK


# ─── PAGE HEADER (template-based) ────────────────────────────────────────────

HEADER_TEMPLATE_PATH = "Headers Template.docx"

def apply_header_placeholders(doc: Document, meta: dict, report_data: dict,
                               cover_db_name: str, client_logo_bytes: Optional[bytes]):
    """Replace {placeholder} tokens in the header loaded from the template.

    Placeholders:
        {logo_client}   → client logo image (or removed if none)
        {CLIENT_NAME }  → client name
        {report_title}  → report title
        {database_name} → cover db name
        {date_created}  → doc date
        {author}        → author name
        {doc_ver}       → doc version
        {doc_rev}       → revision = "0"
    """
    header = doc.sections[0].header
    replacements = {
        "{report_title}":   meta.get("reportTitle") or "",
        "{database_name}":  cover_db_name,
        "{date_created}":   report_data.get("docDate") or "",
        "{author}":         report_data.get("authorName") or "",
        "{doc_ver}":        report_data.get("docVersion") or "",
        "{doc_rev}":        "0",
        "{CLIENT_NAME }":   meta.get("clientName") or "",
    }

    # Walk every <w:t> element in the header and apply plain-text replacements
    for t_elem in header._element.iter(qn("w:t")):
        for placeholder, value in replacements.items():
            if placeholder in (t_elem.text or ""):
                t_elem.text = t_elem.text.replace(placeholder, value)
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Handle {logo_client} — find the run containing it and replace with image
    for para in header.paragraphs:
        for run in para.runs:
            if "{logo_client}" in (run.text or ""):
                run.text = ""  # clear placeholder text
                if client_logo_bytes:
                    try:
                        add_logo_to_run(run, client_logo_bytes,
                                        max_w_inches=1.1, max_h_inches=0.45)
                    except Exception:
                        pass
                break
    # Also check inside table cells in the header
    for tbl in header.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "{logo_client}" in (run.text or ""):
                            run.text = ""
                            if client_logo_bytes:
                                try:
                                    add_logo_to_run(run, client_logo_bytes,
                                                    max_w_inches=1.1, max_h_inches=0.45)
                                except Exception:
                                    pass
                            break


# ─── TABLE OF CONTENTS ────────────────────────────────────────────────────────

def add_toc(doc: Document):
    """Insert a TOC field page. Word fills it in when fields are updated (Ctrl+A → F9)."""
    # Title — plain bold paragraph, NOT a heading, so it doesn't appear inside the TOC itself
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Table of Contents")
    r_title.bold = True
    r_title.font.name = "Cambria"
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = BLUE
    p_title.paragraph_format.space_after = Pt(12)

    # TOC field — one paragraph with four runs: begin / instrText / separate / end
    para = doc.add_paragraph()
    p_xml = para._p

    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    p_xml.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr)
    p_xml.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    p_xml.append(r_sep)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    p_xml.append(r_end)

    # Hint for the user
    p_hint = doc.add_paragraph()
    r_hint = p_hint.add_run(
        "[ Right-click → Update Field, or press Ctrl+A then F9 in Word to populate this TOC ]"
    )
    r_hint.font.name = "Cambria"
    r_hint.font.size = Pt(8)
    r_hint.font.color.rgb = GRAY
    r_hint.italic = True

    add_page_break_para(doc)


# ─── SECTION BUILDERS ─────────────────────────────────────────────────────────

def add_cover_page(doc: Document, meta: dict, report_data: dict, cover_db_name: str,
                   logo_bytes: Optional[bytes], client_logo_bytes: Optional[bytes]):
    if logo_bytes:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            add_logo_to_run(run, logo_bytes, max_w_inches=2.5, max_h_inches=1.2)
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Advance Services Reporting", bold=True, size_pt=18, font="Cambria")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, meta.get("reportTitle") or "", bold=False, size_pt=13, font="Cambria")

    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Report of Healthy Check For", bold=False, size_pt=11, font="Cambria")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Oracle Database {cover_db_name}", bold=True, size_pt=14, font="Cambria")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, report_data.get("period") or "", bold=True, size_pt=13, font="Cambria")

    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Prepared for", bold=False, size_pt=11, font="Cambria")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, meta.get("clientName") or "", bold=True, size_pt=14, font="Cambria")

    if client_logo_bytes:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            add_logo_to_run(run, client_logo_bytes, max_w_inches=2.0, max_h_inches=1.0)
        except Exception:
            pass


def add_info_section(doc: Document, meta: dict, report_data: dict, cover_db_name: str):
    doc.add_heading("1. Information", level=1)
    doc.add_heading("1.1 Document Information", level=2)

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = None
    col_w = [Inches(1.5), Inches(1.8), Inches(1.5), Inches(1.5)]

    rows_data = [
        ("Title", f"{meta.get('reportTitle', '')} - {cover_db_name}", "Doc Version", report_data.get("docVersion") or ""),
        ("Process Owner", meta.get("companyName") or "", "Version Date", report_data.get("docDate") or ""),
        ("Author", report_data.get("authorName") or "", "Date Created", report_data.get("docDate") or ""),
    ]
    for r_idx, (h1, v1, h2, v2) in enumerate(rows_data):
        row = tbl.rows[r_idx]
        hdr_cell(row.cells[0], h1)
        data_cell(row.cells[1], v1)
        hdr_cell(row.cells[2], h2)
        data_cell(row.cells[3], v2)

    doc.add_heading("1.2 Reviewers", level=2)
    reviewers = meta.get("reviewers") or []
    rev_tbl = doc.add_table(rows=1 + len(reviewers), cols=3)
    rev_tbl.style = None
    for i, h in enumerate(["Name", "Position", "Company"]):
        hdr_cell(rev_tbl.rows[0].cells[i], h)
    for r_idx, rev in enumerate(reviewers):
        row = rev_tbl.rows[r_idx + 1]
        data_cell(row.cells[0], rev.get("name") or "")
        data_cell(row.cells[1], rev.get("position") or "")
        data_cell(row.cells[2], rev.get("company") or meta.get("companyName") or "")


def add_exec_summary(doc: Document, meta: dict, report_data: dict,
                     exec_db_text: str, exec_server_text: str):
    add_page_break_para(doc)
    doc.add_heading("2. Executive Summary", level=1)

    p = doc.add_paragraph()
    add_run(p, "ATS conducted a technical assessment and preventive maintenance at ", size_pt=10)
    add_run(p, meta.get("clientName") or "", bold=True, size_pt=10)
    add_run(p, " on the ", size_pt=10)
    add_run(p, exec_db_text, bold=True, size_pt=10)
    add_run(p, " database hosted on ", size_pt=10)
    add_run(p, exec_server_text, bold=True, size_pt=10)
    add_run(p, ". This document recommends changes and identifies specific areas that require more detailed investigation.", size_pt=10)

    p2 = doc.add_paragraph()
    add_run(p2, f"Performance data was collected on {report_data.get('dataCollectionDate') or ''}.", size_pt=10)



def add_system_overview(doc: Document, meta: dict, mode: str, db_info: dict, server_info: dict,
                        single_db: dict, dbs: list):
    add_page_break_para(doc)
    doc.add_heading("3. System Overview", level=1)

    rows_data = [("Client", meta.get("clientName") or "")]
    si = server_info or {}
    di = db_info or {}

    if mode in ("single", "separate"):
        db_name = (single_db or {}).get("databaseName") or meta.get("databaseName") or ""
        rows_data.append(("Database", db_name))
        if di.get("instanceName"):
            rows_data.append(("Instance", di["instanceName"]))
        if di.get("instanceVersion"):
            rows_data.append(("Oracle Version", di["instanceVersion"]))
        if si.get("hostname"):
            host_val = si["hostname"]
            if si.get("ip"):
                host_val += f" ({si['ip']})"
            rows_data.append(("Hostname", host_val))
        if si.get("os"):
            rows_data.append(("Operating System", si["os"]))
        if si.get("kernel"):
            rows_data.append(("Kernel", si["kernel"]))
        cpu_parts = [si.get("cpus"), si.get("systemModel")]
        cpu_str = " — ".join(p for p in cpu_parts if p)
        if cpu_str:
            rows_data.append(("CPU", cpu_str))
        if si.get("totalMemory"):
            rows_data.append(("Total Memory", si["totalMemory"]))
        if di.get("openMode"):
            rows_data.append(("Open Mode", di["openMode"]))
        if di.get("archiveMode"):
            rows_data.append(("Archive Mode", di["archiveMode"]))
        if di.get("characterSet"):
            rows_data.append(("Character Set", di["characterSet"]))
        server_type = (single_db or {}).get("serverType") or ""
        if not si.get("hostname") and server_type:
            rows_data.append(("Server / OS", server_type))
    else:
        for i, db in enumerate(dbs or []):
            label = "Database" if i == 0 else f"Database {i + 1}"
            rows_data.append((label, f"{db.get('databaseName') or '-'} ({db.get('serverType') or '-'})"))

    sys_tbl = doc.add_table(rows=len(rows_data), cols=2)
    sys_tbl.style = None
    for r_idx, (label, value) in enumerate(rows_data):
        hdr_cell(sys_tbl.rows[r_idx].cells[0], label)
        data_cell(sys_tbl.rows[r_idx].cells[1], value)


def add_findings_section(doc: Document, report_data: dict, dbs_data: list, mode: str,
                         cover_db_name: str, logo_bytes, client_logo_bytes):
    doc.add_heading("4. Findings", level=1)
    sections = report_data.get("sections") or []

    if mode == "combined" and dbs_data:
        for sec_idx, sec in enumerate(sections):
            doc.add_heading(f"4.{sec_idx + 1} {sec.get('name', '')}", level=2)
            for db_idx, db_data in enumerate(dbs_data):
                db_sections = db_data.get("sections") or []
                db_sec = db_sections[sec_idx] if sec_idx < len(db_sections) else sec
                db_name = (dbs_data[db_idx].get("databaseName") or
                           f"Database {db_idx + 1}")
                color_rgb = DB_COLORS[db_idx % len(DB_COLORS)]
                _push_sec_content(doc, db_sec, db_name, color_rgb, inline_title=True)
    else:
        for sec_idx, sec in enumerate(sections):
            doc.add_heading(f"4.{sec_idx + 1} {sec.get('name', '')}", level=2)
            _push_sec_content(doc, sec, cover_db_name, (0x0F, 0x4C, 0x75), inline_title=False)


def _push_sec_content(doc: Document, sec: dict, caption_db: str, color_rgb: tuple, inline_title: bool):
    content_type = sec.get("contentType") or ("image" if sec.get("id") == "cpu" else "table")
    columns = sec.get("tableColumns") or []
    rows = sec.get("tableRows") or []

    if content_type in ("table", "table+image") and columns and rows:
        n_cols = len(columns)
        n_rows = len(rows)
        tbl = doc.add_table(rows=(2 if inline_title else 1) + n_rows, cols=n_cols)
        tbl.style = None

        start_row = 0
        if inline_title:
            title_row = tbl.rows[0]
            title_cell = title_row.cells[0]
            # Merge across all columns
            for i in range(1, n_cols):
                title_cell = title_cell.merge(tbl.rows[0].cells[i])
            hdr_cell(title_cell, caption_db, rgb=color_rgb, size_pt=11, bold=True)
            start_row = 1

        col_row = tbl.rows[start_row]
        for i, col_name in enumerate(columns):
            hdr_cell(col_row.cells[i], col_name, rgb=color_rgb)

        for r_idx, row_data in enumerate(rows):
            tbl_row = tbl.rows[start_row + 1 + r_idx]
            for c_idx, cell_val in enumerate(row_data):
                if c_idx < n_cols:
                    data_cell(tbl_row.cells[c_idx], str(cell_val or ""))

        doc.add_paragraph()

    # Images / chart
    if content_type in ("image", "table+image") or (content_type not in ("none", "table")):
        if sec.get("id") == "cpu":
            if inline_title:
                p_title = doc.add_paragraph()
                run = p_title.add_run(caption_db)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*color_rgb)
            image_list = sec.get("chartImageDataList") or []
            if not image_list and sec.get("chartImageData"):
                image_list = [sec["chartImageData"]]
            if image_list:
                for img_b64 in image_list:
                    img_bytes = b64_to_bytes(img_b64)
                    if img_bytes:
                        insert_image(doc, img_bytes, width_inches=5.5)
            else:
                cpu_data = sec.get("cpuData") or {}
                if cpu_data.get("rows"):
                    chart_bytes = render_cpu_chart(cpu_data)
                    if chart_bytes:
                        insert_image(doc, chart_bytes, width_inches=5.5)
        else:
            if inline_title and (sec.get("screenshots") or []):
                p_title = doc.add_paragraph()
                run = p_title.add_run(caption_db)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*color_rgb)
            for ss in (sec.get("screenshots") or []):
                if isinstance(ss, dict):
                    img_bytes = b64_to_bytes(ss.get("data") or "")
                elif isinstance(ss, str):
                    img_bytes = b64_to_bytes(ss)
                else:
                    img_bytes = None
                if img_bytes:
                    insert_image(doc, img_bytes, width_inches=5.0)

    # Recommendation
    p_rec = doc.add_paragraph()
    add_run(p_rec, "Recommendation: ", bold=True, size_pt=10)
    add_run(p_rec, sec.get("recommendation") or "", bold=False, size_pt=10)
    doc.add_paragraph()


def add_summary_section(doc: Document, report_data: dict, sum_db_name: str):
    add_page_break_para(doc)
    doc.add_heading("5. Summary & Recommendation", level=1)

    condition = report_data.get("overallCondition") or "good condition"
    p = doc.add_paragraph()
    add_run(p, "From the general checking of the database, the ", size_pt=10)
    add_run(p, sum_db_name, bold=True, size_pt=10)
    add_run(p, " database is in ", size_pt=10)
    add_run(p, condition, bold=True, size_pt=10)
    add_run(p, " and has acceptable performance.", size_pt=10)

    if report_data.get("summaryText"):
        p2 = doc.add_paragraph()
        add_run(p2, report_data["summaryText"], size_pt=10)


def add_acceptance_section(doc: Document, meta: dict, report_data: dict):
    doc.add_heading("6. Acceptance", level=1)

    p = doc.add_paragraph()
    add_run(p, "I have read, understand and agree to the contents of this document and am authorized to act on behalf of the company.", size_pt=10)

    acc_tbl = doc.add_table(rows=2, cols=2)
    acc_tbl.style = None
    hdr_cell(acc_tbl.rows[0].cells[0], meta.get("clientName") or "Client")
    hdr_cell(acc_tbl.rows[0].cells[1], meta.get("companyName") or "Company")

    data_cell(acc_tbl.rows[1].cells[0], "Authorized Representative:\n\n\nSignature:\n\nDate:")
    data_cell(acc_tbl.rows[1].cells[1],
              f"DBA:\n{report_data.get('authorName') or ''}\n\nSignature:\n\nDate: {report_data.get('docDate') or ''}")
    # Make signature row taller
    set_row_height(acc_tbl.rows[1], 2000)


# ─── MAIN ENTRY POINT ─────────────────────────────────────────────────────────

def build_report(payload: dict) -> bytes:
    """Build .docx from report payload dict. Returns raw bytes."""
    meta = payload.get("meta") or {}
    mode = payload.get("mode") or "single"
    report_data = payload.get("reportData") or {}
    single_db = payload.get("db") or {}
    dbs = payload.get("dbs") or []
    dbs_data = payload.get("dbsData") or []
    server_info = payload.get("serverInfo") or {}
    db_info = payload.get("databaseInfo") or {}
    logo_bytes = b64_to_bytes(payload.get("logoData") or "")
    client_logo_bytes = b64_to_bytes(payload.get("clientLogoData") or "")

    # Resolve display names
    if mode == "single":
        cover_db_name = single_db.get("databaseName") or meta.get("databaseName") or ""
        exec_db_text = cover_db_name
        exec_server_text = single_db.get("serverType") or server_info.get("os") or "the server"
        sum_db_name = cover_db_name
    elif mode == "separate":
        cover_db_name = single_db.get("databaseName") or ""
        exec_db_text = cover_db_name
        exec_server_text = single_db.get("serverType") or server_info.get("os") or "the server"
        sum_db_name = cover_db_name
    else:  # combined
        cover_db_name = ", ".join(d.get("databaseName") or "" for d in dbs if d.get("databaseName"))
        exec_db_text = cover_db_name
        server_types = list(dict.fromkeys(d.get("serverType") or "" for d in dbs if d.get("serverType")))
        exec_server_text = ", ".join(server_types) or "the server"
        sum_db_name = cover_db_name

    # Enrich sections with cpuData if present in payload
    sections = report_data.get("sections") or []
    for sec in sections:
        if sec.get("id") == "cpu" and not sec.get("cpuData"):
            # cpuData may be passed at top level
            cpu_data = payload.get("cpuData")
            if cpu_data:
                sec["cpuData"] = cpu_data

    # Open the header template as the base document so the Sisindokom logo
    # and exact table structure are preserved without any manual XML copying.
    try:
        doc = Document(HEADER_TEMPLATE_PATH)
    except Exception:
        doc = Document()

    apply_doc_styles(doc)

    # Clear the template body (keep sectPr / header relationships)
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    # Page setup — A4 with 2.54 cm margins
    section_obj = doc.sections[0]
    section_obj.page_width = Cm(21)
    section_obj.page_height = Cm(29.7)
    section_obj.left_margin = Cm(2.54)
    section_obj.right_margin = Cm(2.54)
    section_obj.top_margin = Cm(3.0)   # slightly more to leave room for header
    section_obj.bottom_margin = Cm(2.54)
    section_obj.header_distance = Cm(1.27)

    # Replace header placeholders (template logo already embedded)
    apply_header_placeholders(doc, meta, report_data, cover_db_name, client_logo_bytes)

    # Build content
    add_cover_page(doc, meta, report_data, cover_db_name, logo_bytes, client_logo_bytes)
    doc.add_page_break()

    add_toc(doc)
    add_info_section(doc, meta, report_data, cover_db_name)
    add_exec_summary(doc, meta, report_data, exec_db_text, exec_server_text)
    add_system_overview(doc, meta, mode, db_info, server_info, single_db, dbs)
    add_findings_section(doc, report_data, dbs_data, mode, cover_db_name, logo_bytes, client_logo_bytes)
    add_summary_section(doc, report_data, sum_db_name)
    add_acceptance_section(doc, meta, report_data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
